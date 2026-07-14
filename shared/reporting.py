from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from io import BytesIO

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from shared.config import GMV_CURRENCY, seller_label
from shared.metrics import contribution, previous_period, safe_change, summarize_kpis


@dataclass(frozen=True)
class WeeklyReport:
    start: date
    end: date
    previous_start: date
    previous_end: date
    category: str
    current: object
    previous: object
    changes: dict
    top_sellers: pd.DataFrame
    top_categories: pd.DataFrame
    core_judgment: str
    risk_signals: list[str]
    actions: list[str]


def build_weekly_report(current: pd.DataFrame, previous: pd.DataFrame, start: date, end: date, category: str) -> WeeklyReport:
    current_kpi = summarize_kpis(current)
    previous_kpi = summarize_kpis(previous)
    prev_start, prev_end = previous_period(start, end)
    changes = {
        "gmv": safe_change(current_kpi.gmv, previous_kpi.gmv),
        "orders": safe_change(current_kpi.orders, previous_kpi.orders),
        "aov": safe_change(current_kpi.aov, previous_kpi.aov),
        "sellers": safe_change(current_kpi.active_sellers, previous_kpi.active_sellers),
    }
    late_rate = float(current["is_late"].mean()) if not current.empty else 0.0
    low_review_rate = float(current["review_score"].le(2).mean()) if not current.empty else 0.0
    risks: list[str] = []
    actions: list[str] = []
    if changes["gmv"] is not None and changes["gmv"] <= -0.10:
        risks.append(f"GMV较前置期下降{abs(changes['gmv']):.1%}，经营规模需要优先核查。")
        actions.append("拆解订单量与客单价变化，并核查主要商家和品类贡献。")
    if changes["orders"] is not None and changes["orders"] <= -0.10:
        risks.append(f"订单量较前置期下降{abs(changes['orders']):.1%}，是潜在规模压力信号。")
        actions.append("定位订单下降集中的商家、品类和日期，形成优先核查清单。")
    if late_rate >= 0.20:
        risks.append(f"当前期延迟率为{late_rate:.1%}，履约体验需要关注。")
        actions.append("抽查延迟订单及关联商家，区分商家准备与运输阶段差异。")
    if low_review_rate >= 0.20:
        risks.append(f"当前期低评分率为{low_review_rate:.1%}，用户体验存在风险信号。")
        actions.append("复核低评分订单的履约、商品与服务信息，避免直接推断单一待验证原因。")
    if not risks:
        risks.append("当前周期未识别到显著负向异常。")
        actions.append("保持现有监控口径，并在下一个周期复核核心指标和体验护栏。")
    actions.append("后续以GMV、订单量、延迟率和低评分率作为验证指标，不提前声称业务提升。")
    gmv_change = changes["gmv"]
    orders_change = changes["orders"]
    aov_change = changes["aov"]
    if gmv_change is None:
        judgment = "前置期缺少可比基准，当前结果仅用于描述本期经营规模。"
    elif gmv_change < -0.05 and orders_change is not None and aov_change is not None:
        if orders_change < aov_change - 0.03:
            judgment = f"本期GMV较前置期下降{abs(gmv_change):.1%}，订单量收缩是更值得优先核查的变化。"
        elif aov_change < orders_change - 0.03:
            judgment = f"本期GMV较前置期下降{abs(gmv_change):.1%}，客单价下降是更值得优先核查的变化。"
        else:
            judgment = f"本期GMV较前置期下降{abs(gmv_change):.1%}，订单量与客单价均未呈现单一主导驱动。"
    elif gmv_change > 0.05:
        judgment = f"本期GMV较前置期增长{gmv_change:.1%}，建议同时复核增长来源与体验护栏。"
    else:
        judgment = "本期GMV整体相对平稳，当前未识别到单一主导变化。"
    return WeeklyReport(
        start=start,
        end=end,
        previous_start=prev_start.date(),
        previous_end=prev_end.date(),
        category=category,
        current=current_kpi,
        previous=previous_kpi,
        changes=changes,
        top_sellers=contribution(current, "seller_id", limit=5),
        top_categories=contribution(current, "category", limit=5),
        core_judgment=judgment,
        risk_signals=risks[:4],
        actions=list(dict.fromkeys(actions))[:4],
    )


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_width(cell, dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")
    if tc_w.getparent() is None:
        tc_pr.append(tc_w)


def _format_table(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            _set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9)


def _change_text(value: float | None) -> str:
    return "不可比" if value is None else f"{value:+.1%}"


def create_weekly_report_docx(report: WeeklyReport, report_title: str) -> bytes:
    """Create a compact management weekly report using standard_business_brief tokens."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after in [("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6)]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(46, 116, 181)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "电商经营分析中心｜自动经营周报"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(100, 100, 100)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(report_title)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(23)
    subtitle = doc.add_paragraph(f"报告周期：{report.start} 至 {report.end}｜品类：{report.category}")
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(90, 90, 90)

    doc.add_heading("一、管理层摘要", level=1)
    summary = doc.add_paragraph()
    summary.add_run("核心判断：").bold = True
    summary.add_run(report.core_judgment)
    risk = doc.add_paragraph()
    risk.add_run("关键风险信号：").bold = True
    risk.add_run(report.risk_signals[0])
    action = doc.add_paragraph()
    action.add_run("下周期行动：").bold = True
    action.add_run(report.actions[0])

    table = doc.add_table(rows=2, cols=4)
    headers = ["GMV", "订单量", "客单价", "活跃商家"]
    values = [
        f"{report.current.gmv:,.0f}\n{_change_text(report.changes['gmv'])}",
        f"{report.current.orders:,}\n{_change_text(report.changes['orders'])}",
        f"{report.current.aov:,.1f}\n{_change_text(report.changes['aov'])}",
        f"{report.current.active_sellers:,}\n{_change_text(report.changes['sellers'])}",
    ]
    for i, header_text in enumerate(headers):
        table.cell(0, i).text = header_text
        table.cell(1, i).text = values[i]
        _set_cell_shading(table.cell(0, i), "F2F4F7")
        table.cell(0, i).paragraphs[0].runs[0].bold = True
        table.cell(1, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _format_table(table, [2340, 2340, 2340, 2340])

    doc.add_heading("二、经营结构", level=1)
    doc.add_paragraph("以下对象按当前期GMV贡献排序，用于识别主要经营来源，不代表绩效归因。")
    structure = doc.add_table(rows=1, cols=4)
    for i, value in enumerate(["类型", "对象", "GMV", "GMV占比"]):
        structure.cell(0, i).text = value
        structure.cell(0, i).paragraphs[0].runs[0].bold = True
        _set_cell_shading(structure.cell(0, i), "F2F4F7")
    for kind, frame, key in [("商家", report.top_sellers.head(3), "seller_id"), ("品类", report.top_categories.head(3), "category")]:
        for row in frame.itertuples(index=False):
            cells = structure.add_row().cells
            label = seller_label(getattr(row, key)) if key == "seller_id" else str(getattr(row, key))
            for cell, value in zip(cells, [kind, label, f"{row.gmv:,.0f}", f"{row.gmv_share:.1%}"]):
                cell.text = value
    _format_table(structure, [1000, 3560, 2400, 2400])

    doc.add_heading("三、风险信号与优先核查", level=1)
    for signal in report.risk_signals:
        doc.add_paragraph(signal, style="List Bullet")
    doc.add_paragraph("说明：上述内容为数据异常信号或潜在机制，不代表已确认待验证原因。")

    doc.add_heading("四、下周期行动建议", level=1)
    for index, action in enumerate(report.actions, start=1):
        doc.add_paragraph(f"{index}. {action}")

    doc.add_heading("五、口径与证据边界", level=1)
    doc.add_paragraph(
        f"当前期：{report.start}至{report.end}；前置期：{report.previous_start}至{report.previous_end}。"
        "有效订单为已交付订单；GMV为商品金额之和，不含运费，单位BRL。"
        "公开数据不能直接证明库存、曝光、竞争价格、毛利或真实运营动作效果；未经真实试点不声称建议已经带来业务提升。"
    )

    output = BytesIO()
    doc.save(output)
    return output.getvalue()
